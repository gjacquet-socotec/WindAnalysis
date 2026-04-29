"""
Classe de base abstraite pour la génération de rapports Word.

Cette classe définit le Template Method pattern pour la génération de rapports.
Les sous-classes doivent implémenter la stratégie de remplissage des tableaux.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Any
from datetime import datetime
from docx import Document
from docx.shared import Inches
from src.logger_config import get_logger

logger = get_logger(__name__)


class WordPresenter(ABC):
    """
    Classe de base abstraite pour générer des rapports Word dynamiques.

    Cette classe orchestre la génération de documents Word en:
    1. Chargeant un template Word existant
    2. Remplissant les tableaux (stratégie définie par sous-classes)
    3. Remplaçant les métadonnées {{variable}}
    4. Insérant les images de visualisation
    5. Sauvegardant le document généré

    Les sous-classes doivent implémenter _process_tables() avec leur stratégie spécifique.
    """

    def __init__(self, template_path: str, output_path: str):
        """
        Initialise le présentateur Word.

        Args:
            template_path: Chemin vers le template Word (.docx)
            output_path: Chemin où sauvegarder le rapport généré

        Raises:
            FileNotFoundError: Si le template n'existe pas
        """
        self.template_path = Path(template_path)
        self.output_path = Path(output_path)

        if not self.template_path.exists():
            raise FileNotFoundError(
                f"Template not found: {self.template_path}. "
                f"Please provide a valid template path."
            )

        logger.info(
            f"WordPresenter initialized with template: {self.template_path.name}"
        )

    def render_report(
        self, context: Dict[str, Any], metadata: Dict[str, Any] = None
    ) -> None:
        """
        Génère un rapport Word en remplissant le template (Template Method Pattern).

        Args:
            context: Dictionnaire avec toutes les données pour les tableaux
                    Format: {
                        'table_name': [{'col1': val1, 'col2': val2}, ...],
                        'chart_paths': {'chart_name': '/path/to/chart.png', ...},
                        ...
                    }
            metadata: Métadonnées additionnelles (dates, turbines, critères)

        Raises:
            Exception: Si la génération du document échoue
        """
        logger.info(f"Rendering Word report from template: {self.template_path}")

        try:
            # Charger le template avec python-docx
            doc = Document(self.template_path)

            # Préparer le contexte complet (métadonnées + données)
            full_context = self._prepare_context(context, metadata)

            # Remplir les tableaux (stratégie spécifique à chaque workflow)
            self._process_tables(doc, full_context)

            # Remplacer les balises de métadonnées dans les paragraphes
            self._replace_metadata_tags(doc, full_context)

            # Insérer les images de visualisation
            self._insert_images(doc, full_context)

            # Créer le dossier output si nécessaire
            self.output_path.parent.mkdir(parents=True, exist_ok=True)

            # Sauvegarder le document
            doc.save(self.output_path)

            logger.info(f"✅ Report successfully saved to: {self.output_path}")

        except Exception as e:
            logger.error(f"❌ Failed to generate Word report: {e}")
            raise

    @abstractmethod
    def _process_tables(self, doc: Document, context: Dict[str, Any]) -> None:
        """
        Remplit les tableaux du document avec les données du contexte.

        Cette méthode abstraite doit être implémentée par les sous-classes
        avec leur stratégie spécifique de remplissage (par index, par marqueur, etc.).

        Args:
            doc: Document Word chargé
            context: Contexte complet avec données et métadonnées
        """
        pass

    def _prepare_context(
        self, context: Dict[str, Any], metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """
        Prépare le contexte complet pour le template avec métadonnées.

        Args:
            context: Données de tables et chemins d'images
            metadata: Métadonnées (dates, turbines, critères)

        Returns:
            Contexte complet fusionnant context + metadata + generation_date
        """
        full_context = context.copy()

        # Ajouter la date de génération
        full_context["generation_date"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Ajouter les métadonnées si fournies
        if metadata:
            # Période de test/analyse
            if "test_start" in metadata:
                test_start = metadata["test_start"]
                if hasattr(test_start, "strftime"):
                    test_start = test_start.strftime("%Y-%m-%d %H:%M:%S")
                full_context["test_start"] = test_start

            if "test_end" in metadata:
                test_end = metadata["test_end"]
                if hasattr(test_end, "strftime"):
                    test_end = test_end.strftime("%Y-%m-%d %H:%M:%S")
                full_context["test_end"] = test_end

            # Pour SCADA: analysis_start et analysis_end
            if "analysis_start" in metadata:
                analysis_start = metadata["analysis_start"]
                if hasattr(analysis_start, "strftime"):
                    analysis_start = analysis_start.strftime("%Y-%m-%d %H:%M:%S")
                full_context["analysis_start"] = analysis_start

            if "analysis_end" in metadata:
                analysis_end = metadata["analysis_end"]
                if hasattr(analysis_end, "strftime"):
                    analysis_end = analysis_end.strftime("%Y-%m-%d %H:%M:%S")
                full_context["analysis_end"] = analysis_end

            # Liste des turbines analysées
            if "turbines" in metadata:
                turbines = metadata["turbines"]
                if isinstance(turbines, list):
                    turbine_list = ", ".join(turbines)
                    full_context["turbine_list"] = turbine_list
                    full_context["turbine_count"] = len(turbines)

            # Autres métadonnées (parc, critères, etc.)
            full_context.update(metadata)

        logger.info(
            f"Context prepared with {len(full_context)} keys "
            f"(generation_date: {full_context['generation_date']})"
        )

        return full_context

    def _replace_metadata_tags(self, doc: Document, context: Dict[str, Any]) -> None:
        """
        Remplace les balises de métadonnées {{variable}} dans les paragraphes.

        Args:
            doc: Document Word
            context: Contexte avec les métadonnées et critères
        """
        logger.info("Replacing metadata tags...")

        replacements = {
            # Métadonnées générales
            "{{generation_date}}": context.get("generation_date", ""),
            "{{test_start}}": context.get("test_start", ""),
            "{{test_end}}": context.get("test_end", ""),
            "{{analysis_start}}": context.get("analysis_start", ""),
            "{{analysis_end}}": context.get("analysis_end", ""),
            "{{turbine_list}}": context.get("turbine_list", ""),
            "{{turbine_count}}": str(context.get("turbine_count", "")),
            # Informations sur le parc
            "{{park_name}}": context.get("park_name", ""),
            "{{Nom_du_parc}}": context.get("park_name", ""),
            "{{Nombre_WTG}}": str(context.get("turbine_count", "")),
            "{{model_wtg}}": context.get("model_wtg", ""),
            "{{Modèle_turbine}}": context.get("model_wtg", ""),
            "{{nominal_power}}": context.get("nominal_power", ""),
            "{{Puissance_nominale}}": context.get("nominal_power", ""),
            # Valeurs des critères de validation (RunTest)
            "{{consecutive_hours_h}}": str(context.get("consecutive_hours_h", "120")),
            "{{cut_in_to_cut_out_h}}": str(context.get("cut_in_to_cut_out_h", "72")),
            "{{cut_in_v_min}}": str(context.get("cut_in_v_min", "3")),
            "{{cut_in_v_max}}": str(context.get("cut_in_v_max", "25")),
            "{{nominal_power_h}}": str(context.get("nominal_power_h", "3")),
            "{{ nominal_power_h}}": str(context.get("nominal_power_h", "3")),
            "{{nominal_power_pct}}": str(context.get("nominal_power_pct", "97")),
            "{{nominal_power_pct }}": str(context.get("nominal_power_pct", "97")),
            "{{ nominal_power_pct}}": str(context.get("nominal_power_pct", "97")),
            "{{local_restarts_max}}": str(context.get("local_restarts_max", "3")),
            "{{local_restarts_max }}": str(context.get("local_restarts_max", "3")),
            "{{availability_min_pct}}": str(context.get("availability_min_pct", "92")),
            "{{availability_min_pct }}": str(context.get("availability_min_pct", "92")),
        }

        # Parcourir tous les paragraphes
        for para in doc.paragraphs:
            full_text = para.text
            modified = False

            for old_tag, new_value in replacements.items():
                if old_tag in full_text:
                    full_text = full_text.replace(old_tag, str(new_value))
                    modified = True

            # Si le texte a été modifié, reconstruire le paragraphe
            if modified:
                # Sauvegarder le formatage du premier run
                first_run_format = None
                if para.runs:
                    first_run = para.runs[0]
                    first_run_format = {
                        "bold": first_run.bold,
                        "italic": first_run.italic,
                        "underline": first_run.underline,
                        "font_name": first_run.font.name,
                        "font_size": first_run.font.size,
                    }

                # Supprimer tous les runs
                for run in para.runs:
                    run.text = ""

                # Ajouter le nouveau texte dans le premier run
                if para.runs:
                    para.runs[0].text = full_text
                    # Restaurer le formatage
                    if first_run_format:
                        para.runs[0].bold = first_run_format["bold"]
                        para.runs[0].italic = first_run_format["italic"]
                        para.runs[0].underline = first_run_format["underline"]
                        if first_run_format["font_name"]:
                            para.runs[0].font.name = first_run_format["font_name"]
                        if first_run_format["font_size"]:
                            para.runs[0].font.size = first_run_format["font_size"]
                else:
                    para.add_run(full_text)

        logger.info("✅ Metadata tags replaced")

    def _insert_images(self, doc: Document, context: Dict[str, Any]) -> None:
        """
        Insère les images de visualisation à la place des balises dans le document.

        Args:
            doc: Document Word
            context: Contexte contenant les chemins des images dans 'chart_paths'
        """
        logger.info("Inserting visualization images...")

        # Récupérer les chemins des images depuis le contexte
        chart_paths = context.get("chart_paths", {})

        if not chart_paths:
            logger.warning("⚠️ No chart paths found in context")
            return

        # Mapping des balises vers les noms de charts (RunTest)
        runtest_image_tags = {
            "{{wind_rose_visualizer}}": "wind_rose_chart",
            "{{wind_bin_visualizer}}": "wind_histogram_chart",
            "{{power_curve_visualizer}}": "power_curve_chart",
            "{{cut_in_cut_out_timeline_visualizer}}": "cutin_cutout_timeline_chart",
            "{{heatmap_chart}}": "heatmap_chart",
        }

        # Mapping des balises SCADA
        scada_image_tags = {
            "{{eba_cut_in_cut_out_chart}}": "eba_cut_in_cut_out_chart",
            "{{eba_manifacturer_chart}}": "eba_manifacturer_chart",
            "{{eba_loss_chart}}": "eba_loss_chart",
            "{{top_error_code_frequency}}": "top_error_code_frequency",
            "{{treemap_error_code}}": "treemap_error_code",
            "{{data_availability_chart}}": "data_availability_chart",
            "{{wind_direction_calibration}}": "wind_direction_calibration",
            "{{power_rose_chart}}": "power_rose_chart",
            "{{wind_rose_chart}}": "wind_rose_chart",
            "{{rpm_chart}}": "rpm_chart",
        }

        # Fusionner les deux mappings
        image_tags = {**runtest_image_tags, **scada_image_tags}

        # Parcourir tous les paragraphes pour trouver les balises
        for para in doc.paragraphs:
            for tag, chart_name in image_tags.items():
                if tag in para.text:
                    # Vérifier si l'image existe
                    if chart_name not in chart_paths:
                        logger.warning(f"⚠️ Image not found for {chart_name}")
                        continue

                    image_path = chart_paths[chart_name]
                    if not Path(image_path).exists():
                        logger.warning(f"⚠️ Image file not found: {image_path}")
                        continue

                    # Supprimer le texte de la balise
                    para.text = ""

                    # Insérer l'image
                    run = para.add_run()
                    try:
                        run.add_picture(str(image_path), width=Inches(6.0))
                        logger.info(f"✅ Image inserted: {chart_name}")
                    except Exception as e:
                        logger.error(f"❌ Error inserting {chart_name}: {e}")

        logger.info("✅ Images inserted in document")
