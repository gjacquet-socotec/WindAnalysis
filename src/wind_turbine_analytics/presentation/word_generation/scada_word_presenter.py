"""
Présentateur Word spécifique au workflow SCADA.

Utilise une stratégie de remplissage des tableaux par marqueurs invisibles.
Le template contient des marqueurs [TABLE:xxx] avant chaque tableau.
"""

from typing import Dict, Any, List
from docx import Document
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

import re
from .word_presenter import WordPresenter
from src.logger_config import get_logger

logger = get_logger(__name__)


class ScadaWordPresenter(WordPresenter):
    """
    Présentateur Word pour les rapports SCADA.

    Stratégie de remplissage des tableaux:
    - Template contient des marqueurs invisibles [TABLE:xxx] avant chaque tableau
    - On cherche chaque marqueur et on trouve le tableau qui suit
    - Remplissage dynamique basé sur le nom du marqueur
    - Ordre flexible, nombre de tables variable
    - Supporte l'ajout/suppression de tables sans changer le code

    Marqueurs attendus dans le template SCADA:
    - [TABLE:CSV_FILES_TABLE]
    - [TABLE:SCADA_SUMMARY_TABLE]
    - [TABLE:EBA_CUT_IN_CUT_OUT_TABLE]
    - [TABLE:EBA_MANUFACTURER_TABLE]
    - [TABLE:EBA_LOSS_TABLE]
    - [TABLE:ERROR_CODES_TABLE]
    - [TABLE:DATA_AVAILABILITY_TABLE]
    - [TABLE:WIND_CALIBRATION_TABLE]
    - [TABLE:TIP_SPEED_RATIO_TABLE]
    - [TABLE:NORMATIVE_YIELD_TABLE]
    - [TABLE:CODE_ERROR_PARETO_FREQUENCY]
    """

    def _set_table_borders(self, table):
        """Bordures fines et élégantes (couleur gris bleu)."""
        tbl = table._element
        tblPr = tbl.xpath("w:tblPr")[0]
        tblBorders = OxmlElement("w:tblBorders")

        # Couleur des bordures : un bleu/gris discret
        border_color = "A6ACAF"
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")  # Très fin
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), border_color)
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def _set_cell_shading(self, cell, color):
        """Applique une couleur de fond à une cellule (Ex: '2F5597' pour bleu foncé)."""
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

    def _process_tables(self, doc: Document, context: Dict[str, Any]) -> None:
        logger.info("Replacing markers with dynamic tables...")

        table_regex = re.compile(r"\[TABLE:\s*(.*?)\s*\]", re.IGNORECASE)

        # On parcourt les paragraphes du document
        for para in list(
            doc.paragraphs
        ):  # On utilise list() pour pouvoir modifier le doc en bouclant
            match = table_regex.search(para.text)

            if match:
                tag_content = match.group(1).strip().upper()
                logger.info(f"Marker found: {tag_content}")

                # Chercher les données
                data = None
                for k, v in context.items():
                    if k.strip().upper() == tag_content:
                        data = v
                        break

                if data and isinstance(data, list) and len(data) > 0:
                    # 1. Créer le tableau à l'endroit du paragraphe
                    new_table = self._insert_table_at_paragraph(doc, para, data)

                    # 2. Remplir le tableau
                    self._populate_table_dynamically(new_table, data)

                    # 3. Supprimer le paragraphe qui contenait le tag [TABLE:XXX]
                    self._remove_paragraph(para)

                    logger.info(
                        f"✅ Table '{tag_content}' created and replaced marker."
                    )
                else:
                    logger.warning(
                        f"⚠️ No data or empty list for '{tag_content}'. Marker left as is."
                    )

    def _insert_table_at_paragraph(
        self, doc: Document, para, data: List[Dict[str, Any]]
    ) -> Table:
        num_cols = len(data[0].keys())
        num_rows = 1

        table = doc.add_table(rows=num_rows, cols=num_cols)

        # On applique nos bordures manuelles au lieu du style qui plante
        self._set_table_borders(table)

        # Déplacer le tableau après le marqueur
        para._element.addnext(table._element)
        return table

    def _remove_paragraph(self, paragraph):
        """Supprime proprement un paragraphe du document."""
        p = paragraph._element
        p.getparent().remove(p)

    def _populate_table_dynamically(
        self, table: Table, data: List[Dict[str, Any]]
    ) -> None:
        headers = list(data[0].keys())

        # --- STYLE DE L'ENTÊTE (Bleu foncé, Texte blanc) ---
        header_bg_color = "2F5597"  # Bleu pro
        for i, header_text in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = str(header_text)
            # Fond bleu foncé
            self._set_cell_shading(cell, header_bg_color)

            # Texte blanc et gras
            paragraph = cell.paragraphs[0]
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc

        # --- STYLE DES LIGNES (Alternance de blanc et bleu très clair) ---
        alternate_row_color = "EBF1DE"  # Ou "D9E1F2" pour un bleu très léger
        for row_idx, row_dict in enumerate(data):
            new_row = table.add_row()

            # Déterminer si c'est une ligne alternée
            is_alternate = row_idx % 2 == 0  # Changez à 0 ou 1 selon votre préférence

            for col_idx, header_name in enumerate(headers):
                cell = new_row.cells[col_idx]
                val = row_dict.get(header_name, "")

                # Injection de la valeur
                cell.text = f"{val:.2f}" if isinstance(val, float) else str(val)

                # Appliquer la couleur de fond alternée
                if is_alternate:
                    self._set_cell_shading(cell, "D9E1F2")  # Bleu très clair

                # Optionnel : alignement à gauche pour le texte, droite pour les chiffres
                if isinstance(val, (int, float)):
                    cell.paragraphs[0].alignment = (
                        2  # Droite (WD_ALIGN_PARAGRAPH.RIGHT)
                    )
