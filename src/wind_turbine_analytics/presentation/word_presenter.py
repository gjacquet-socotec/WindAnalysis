from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from src.logger_config import get_logger

logger = get_logger(__name__)


class WordPresenter:
    """
    Génère des rapports Word dynamiques avec python-docx.

    Cette classe orchestre la génération de documents Word en:
    1. Chargeant un template Word existant (ou en créant un nouveau)
    2. Identifiant les tableaux à remplir (par marqueurs)
    3. Remplissant les tableaux dynamiquement avec les données
    4. Ajoutant les métadonnées dans le document
    5. Sauvegardant le document généré

    Approche: Utilise python-docx directement (pas docxtpl) pour avoir
    un contrôle total sur la génération des tableaux dynamiques.
    """

    def __init__(
        self,
        template_path: str,
        output_path: str,
        auto_create_template: bool = True
    ):
        """
        Initialise le présentateur Word.

        Args:
            template_path: Chemin vers le template Word (.docx)
            output_path: Chemin où sauvegarder le rapport généré
            auto_create_template: Si True, crée automatiquement un template optimisé
                                  s'il n'existe pas déjà

        Raises:
            FileNotFoundError: Si le template n'existe pas et auto_create_template=False
        """
        self.template_path = Path(template_path)
        self.output_path = Path(output_path)

        # Vérifier si un template optimisé existe déjà
        new_template_path = self.template_path.parent / f"{self.template_path.stem}_new{self.template_path.suffix}"

        if new_template_path.exists():
            # Utiliser le template optimisé existant
            logger.info(f"Utilisation du template optimisé existant: {new_template_path.name}")
            self.template_path = new_template_path
        elif self.template_path.exists() and auto_create_template:
            # Créer le template optimisé UNE SEULE FOIS
            logger.info(f"Template optimisé introuvable, création à partir de: {self.template_path.name}")
            self._create_new_template()
            # Utiliser le nouveau template créé
            self.template_path = new_template_path
        elif not self.template_path.exists():
            raise FileNotFoundError(
                f"Template not found: {self.template_path}. "
                f"Please provide a valid template path."
            )

        logger.info(f"WordPresenter prêt avec: {self.template_path.name}")

    def render_report(
        self, context: Dict[str, Any], metadata: Dict[str, Any] = None
    ) -> None:
        """
        Génère un rapport Word en remplissant les tableaux dynamiquement.

        Args:
            context: Dictionnaire avec toutes les données pour les tableaux
                    Format: {
                        'table_name': [{'col1': val1, 'col2': val2}, ...],
                        ...
                    }
            metadata: Métadonnées additionnelles (date, période, turbines)

        Raises:
            Exception: Si la génération du document échoue
        """
        logger.info(f"Rendering Word report from template: {self.template_path}")

        try:
            # Charger le template avec python-docx
            doc = Document(self.template_path)

            # Préparer le contexte complet
            full_context = self._prepare_context(context, metadata)

            # Remplir les tableaux dynamiquement
            self._fill_tables(doc, full_context)

            # Remplacer les balises de métadonnées dans les paragraphes
            self._replace_metadata_tags(doc, full_context)

            # Créer le dossier output si nécessaire
            self.output_path.parent.mkdir(parents=True, exist_ok=True)

            # Sauvegarder le document
            doc.save(self.output_path)

            logger.info(f"✅ Report successfully saved to: {self.output_path}")

        except Exception as e:
            logger.error(f"❌ Failed to generate Word report: {e}")
            raise

    def _OLD_adapt_template_for_docxtpl(self) -> None:
        """
        Adapte automatiquement le template Word avec la syntaxe docxtpl.

        Cette méthode:
        1. Charge le template Word avec python-docx
        2. Identifie les tableaux à adapter
        3. Remplace les anciennes balises par la syntaxe docxtpl avec boucles
        4. Sauvegarde une version adaptée temporaire

        Le template adapté remplace l'original pour la génération.

        Mapping des tableaux (indices python-docx):
        - Tableau 0: Historique des révisions (SKIP)
        - Tableau 1: Périodes de Run Test (SKIP - pas encore implémenté)
        - Tableau 2: Heures consécutives (critère 1)
        - Tableau 3: Cut-in/Cut-out (critère 2)
        - Tableau 4: Puissance nominale - valeurs (critère 3a)
        - Tableau 5: Puissance nominale - durée (critère 3b)
        - Tableau 6: Autonomie (critère 4)
        - Tableau 7: Disponibilité (critère 5)
        - Tableau 8-11: Graphiques (SKIP)

        Note: Il n'y a PAS de tableau récapitulatif global dans le template actuel.
              Le tableau summary_table sera ignoré pour l'instant.
        """
        logger.info("Adaptation automatique du template pour syntaxe docxtpl...")

        try:
            # Charger avec python-docx pour modification
            doc = Document(self.template_path)

            logger.info(f"Template chargé: {len(doc.tables)} tableaux trouvés")

            # Adapter les tableaux avec les BONS indices
            if len(doc.tables) >= 8:
                self._adapt_table_consecutive_hours(doc.tables[2])  # Critère 1
                self._adapt_table_cut_in_cut_out(doc.tables[3])  # Critère 2
                self._adapt_table_nominal_power_values(doc.tables[4])  # Critère 3a
                self._adapt_table_nominal_power_duration(doc.tables[5])  # Critère 3b
                self._adapt_table_autonomous_operation(doc.tables[6])  # Critère 4
                self._adapt_table_availability(doc.tables[7])  # Critère 5
                logger.info("✅ 6 tableaux adaptés")
            else:
                logger.warning(
                    f"Template n'a que {len(doc.tables)} tableaux, "
                    f"besoin d'au moins 8 pour adaptation complète"
                )

            # Sauvegarder le template adapté (temporaire)
            adapted_path = self.template_path.parent / f"{self.template_path.stem}_adapted{self.template_path.suffix}"
            doc.save(adapted_path)

            # Utiliser le template adapté pour la génération
            self.template_path = adapted_path

            logger.info(f"✅ Template adapté sauvegardé: {adapted_path}")

        except Exception as e:
            logger.warning(
                f"⚠️ Échec de l'adaptation automatique du template: {e}. "
                f"Utilisation du template original."
            )

    def _adapt_table_summary(self, table) -> None:
        """Adapte le tableau récapitulatif de synthèse."""
        if len(table.rows) < 2:
            return

        # Headers
        headers = ["WTG", "Criterion 1\n(≥120h)", "Criterion 2\n(≥72h)",
                   "Criterion 3\n(≥3h @97%)", "Criterion 4\n(≤3 restarts)",
                   "Criterion 5\n(≥92%)", "Overall"]

        if len(table.columns) >= 7:
            row_1 = table.rows[0]
            for col_idx, header in enumerate(headers):
                if col_idx < len(row_1.cells):
                    row_1.cells[col_idx].text = header

            # Ligne 2: Syntaxe docxtpl
            row_2 = table.rows[1]
            loop_syntax = (
                "{%tr for item in summary_table%}{{item.wtg}}{%tc%}"
                "{{item.consecutive_hours}}{%tc%}{{item.cut_in_cut_out}}{%tc%}"
                "{{item.nominal_power}}{%tc%}{{item.autonomous_operation}}{%tc%}"
                "{{item.availability}}{%tc%}{{item.overall}}{%endtr%}"
            )
            row_2.cells[0].text = loop_syntax
            for col_idx in range(1, len(row_2.cells)):
                row_2.cells[col_idx].text = ""

    def _adapt_table_consecutive_hours(self, table) -> None:
        """
        Adapte le tableau des heures consécutives.

        Syntaxe docxtpl pour tableaux:
        - Ligne de données : mettre {%tr for item in table%} dans la PREMIÈRE cellule
        - Autres cellules : {{item.field}}
        - Dernière cellule : {{item.field}}{%endtr%}
        """
        if len(table.rows) < 2 or len(table.columns) < 3:
            return

        headers = ["WTG", "Data hours [h]", "Criterion (≥120h)"]
        row_1 = table.rows[0]
        for col_idx, header in enumerate(headers):
            if col_idx < len(row_1.cells):
                row_1.cells[col_idx].text = header

        row_2 = table.rows[1]
        # Syntaxe docxtpl: répartir sur les cellules
        row_2.cells[0].text = "{%tr for item in consecutive_hours_table%}{{item.wtg}}"
        row_2.cells[1].text = "{{item.data_hours}}"
        row_2.cells[2].text = "{{item.criterion_met}}{%endtr%}"

    def _adapt_table_cut_in_cut_out(self, table) -> None:
        """Adapte le tableau cut-in à cut-out."""
        if len(table.rows) < 2 or len(table.columns) < 3:
            return

        headers = ["WTG", "Data hours [h]", "Criterion (≥72h)"]
        row_1 = table.rows[0]
        for col_idx, header in enumerate(headers):
            if col_idx < len(row_1.cells):
                row_1.cells[col_idx].text = header

        row_2 = table.rows[1]
        row_2.cells[0].text = "{%tr for item in cut_in_cut_out_table%}{{item.wtg}}"
        row_2.cells[1].text = "{{item.data_hours}}"
        row_2.cells[2].text = "{{item.criterion_met}}{%endtr%}"

    def _adapt_table_nominal_power_values(self, table) -> None:
        """Adapte le tableau puissance nominale - valeurs."""
        if len(table.rows) < 2 or len(table.columns) < 4:
            return

        headers = ["WTG", "Max Wind Speed [m/s]", "Max Power [kW]", "Status"]
        row_1 = table.rows[0]
        for col_idx, header in enumerate(headers):
            if col_idx < len(row_1.cells):
                row_1.cells[col_idx].text = header

        row_2 = table.rows[1]
        row_2.cells[0].text = "{%tr for item in nominal_power_values_table%}{{item.wtg}}"
        row_2.cells[1].text = "{{item.max_wind_speed_ms}}"
        row_2.cells[2].text = "{{item.max_power_kw}}"
        row_2.cells[3].text = "{{item.status}}{%endtr%}"

    def _adapt_table_nominal_power_duration(self, table) -> None:
        """Adapte le tableau puissance nominale - durée."""
        if len(table.rows) < 2 or len(table.columns) < 4:
            return

        headers = ["WTG", "Duration [h]", "Period", "Status"]
        row_1 = table.rows[0]
        for col_idx, header in enumerate(headers):
            if col_idx < len(row_1.cells):
                row_1.cells[col_idx].text = header

        row_2 = table.rows[1]
        row_2.cells[0].text = "{%tr for item in nominal_power_duration_table%}{{item.wtg}}"
        row_2.cells[1].text = "{{item.duration_hours}}"
        row_2.cells[2].text = "{{item.start}}"
        row_2.cells[3].text = "{{item.status}}{%endtr%}"

        logger.info(f"Tableau puissance nominale - durée adapté: {len(headers)} colonnes")

    def _adapt_table_autonomous_operation(self, table) -> None:
        """Adapte le tableau autonomie d'exploitation."""
        if len(table.rows) < 2 or len(table.columns) < 3:
            return

        headers = ["WTG", "Local Acknowledgements / Restarts", "Criterion (≤3)"]
        row_1 = table.rows[0]
        for col_idx, header in enumerate(headers):
            if col_idx < len(row_1.cells):
                row_1.cells[col_idx].text = header

        row_2 = table.rows[1]
        row_2.cells[0].text = "{%tr for item in autonomous_operation_table%}{{item.wtg}}"
        row_2.cells[1].text = "{{item.manual_restart_count}}"
        row_2.cells[2].text = "{{item.criterion_met}}{%endtr%}"

    def _adapt_table_availability(self, table) -> None:
        """Adapte le tableau disponibilité."""
        if len(table.rows) < 2 or len(table.columns) < 5:
            return

        headers = ["WTG", "Test Duration [h]", "Unauthorized Downtime [h]",
                   "Availability (%)", "Criterion (≥92%)"]
        row_1 = table.rows[0]
        for col_idx, header in enumerate(headers):
            if col_idx < len(row_1.cells):
                row_1.cells[col_idx].text = header

        row_2 = table.rows[1]
        row_2.cells[0].text = "{%tr for item in availability_table%}{{item.wtg}}"
        row_2.cells[1].text = "{{item.wtg_ok_hours}}"
        row_2.cells[2].text = "{{item.warning_hours}}"
        row_2.cells[3].text = "{{item.availability_percent}}"
        row_2.cells[4].text = "{{item.criterion_met}}{%endtr%}"

    def _prepare_context(
        self, context: Dict[str, Any], metadata: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """
        Prépare le contexte complet pour le template avec métadonnées.

        Args:
            context: Données de tables
            metadata: Métadonnées (test_start, test_end, turbines)

        Returns:
            Contexte complet pour docxtpl
        """
        full_context = context.copy()

        # Ajouter la date de génération
        full_context["generation_date"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Ajouter les métadonnées si fournies
        if metadata:
            # Période de test
            if "test_start" in metadata:
                test_start = metadata["test_start"]
                # Formater si c'est un datetime
                if hasattr(test_start, "strftime"):
                    test_start = test_start.strftime("%Y-%m-%d %H:%M:%S")
                full_context["test_start"] = test_start

            if "test_end" in metadata:
                test_end = metadata["test_end"]
                # Formater si c'est un datetime
                if hasattr(test_end, "strftime"):
                    test_end = test_end.strftime("%Y-%m-%d %H:%M:%S")
                full_context["test_end"] = test_end

            # Liste des turbines analysées
            if "turbines" in metadata:
                turbines = metadata["turbines"]
                if isinstance(turbines, list):
                    turbine_list = ", ".join(turbines)
                    full_context["turbine_list"] = turbine_list
                    full_context["turbine_count"] = len(turbines)

            # Autres métadonnées
            full_context.update(metadata)

        logger.info(
            f"Context prepared with {len(full_context)} keys "
            f"(generation_date: {full_context['generation_date']})"
        )

        return full_context

    def _create_new_template(self) -> None:
        """
        Crée un nouveau template optimisé à partir du template original.

        Cette méthode:
        1. Copie tout le contenu du template original
        2. Identifie et remplace les tableaux par des versions propres (en-têtes seulement)
        3. Ajoute le tableau récapitulatif global
        4. Marque les tableaux avec des identifiants pour remplissage dynamique
        """
        logger.info("Création d'un nouveau template optimisé...")

        original_path = self.template_path
        new_path = self.template_path.parent / f"{self.template_path.stem}_new{self.template_path.suffix}"

        if not original_path.exists():
            logger.error(f"Template original introuvable: {original_path}")
            return

        # Charger le template original
        original_doc = Document(original_path)
        new_doc = Document()

        logger.info(f"Template original: {len(original_doc.paragraphs)} paragraphes, "
                    f"{len(original_doc.tables)} tableaux")

        # Copier le contenu en identifiant les tableaux
        table_counter = 0
        for element in original_doc.element.body:
            # Si c'est un paragraphe, le copier
            if element.tag.endswith('p'):
                for para in original_doc.paragraphs:
                    if para._element == element:
                        self._copy_paragraph(para, new_doc)
                        break

            # Si c'est un tableau, le remplacer par une version propre
            elif element.tag.endswith('tbl'):
                table_counter += 1

                # Identifier et créer le bon tableau
                if table_counter == 2:
                    # NOUVEAU: Tableau récapitulatif global
                    new_doc.add_heading('Synthèse des résultats des Run Tests', level=2)
                    self._create_table_header(new_doc, 'summary_table', [
                        'WTG', 'Criterion 1\n(≥120h)', 'Criterion 2\n(≥72h)',
                        'Criterion 3\n(≥3h @97%)', 'Criterion 4\n(≤3 restarts)',
                        'Criterion 5\n(≥92%)', 'Overall'
                    ])
                elif table_counter == 3:
                    new_doc.add_heading('Durée cumulée de fonctionnement', level=2)
                    self._create_table_header(new_doc, 'consecutive_hours_table',
                                              ['WTG', 'Data hours [h]', 'Criterion (≥120h)'])
                elif table_counter == 4:
                    new_doc.add_heading('Fonctionnement dans la plage de vent', level=2)
                    self._create_table_header(new_doc, 'cut_in_cut_out_table',
                                              ['WTG', 'Data hours [h]', 'Criterion (≥72h)'])
                elif table_counter == 5:
                    new_doc.add_heading('Atteinte de la puissance nominale', level=2)
                    self._create_table_header(new_doc, 'nominal_power_values_table',
                                              ['WTG', 'Max Power [kW]', 'Max Wind Speed [m/s]', 'Status'])
                elif table_counter == 6:
                    self._create_table_header(new_doc, 'nominal_power_duration_table',
                                              ['WTG', 'Duration [h]', 'Start', 'End', 'Status'])
                elif table_counter == 7:
                    new_doc.add_heading('Fonctionnement en mode automatique', level=2)
                    self._create_table_header(new_doc, 'autonomous_operation_table',
                                              ['WTG', 'Local Acknowledgements / Restarts', 'Criterion (≤3)'])
                elif table_counter == 8:
                    new_doc.add_heading('Disponibilité des turbines', level=2)
                    self._create_table_header(new_doc, 'availability_table',
                                              ['WTG', 'Availability (%)', 'WTG OK [h]',
                                               'Warning [h]', 'Criterion (≥92%)'])

        # Sauvegarder le nouveau template
        new_path.parent.mkdir(parents=True, exist_ok=True)
        new_doc.save(new_path)

        logger.info(f"✅ Nouveau template créé: {new_path}")
        logger.info(f"Taille: {new_path.stat().st_size / 1024:.2f} KB")

        # Utiliser le nouveau template
        self.template_path = new_path

    def _copy_paragraph(self, source_para, target_doc):
        """Copie un paragraphe avec son style et formatage."""
        new_para = target_doc.add_paragraph()

        # Copier le style
        if source_para.style:
            try:
                new_para.style = source_para.style
            except:
                pass

        # Copier les runs (portions de texte)
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name

        # Copier l'alignement
        new_para.alignment = source_para.alignment

        return new_para

    def _create_table_header(self, doc, table_name: str, headers: List[str]):
        """
        Crée un tableau avec seulement les en-têtes.

        Args:
            doc: Document Word
            table_name: Nom du tableau (pour identification)
            headers: Liste des en-têtes de colonnes
        """
        # Ajouter un marqueur invisible pour identifier le tableau
        marker_para = doc.add_paragraph(f'[TABLE:{table_name.upper()}]')
        marker_para.runs[0].font.size = Pt(1)  # Très petit
        marker_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)  # Blanc

        # Créer le tableau avec en-têtes
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = header
            # Centrer et mettre en gras
            for paragraph in hdr_cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True

        return table

    def _fill_tables(self, doc: Document, context: Dict[str, Any]) -> None:
        """
        Remplit les tableaux du document avec les données du contexte.

        Cette méthode parcourt tous les paragraphes pour trouver les marqueurs
        [TABLE:xxx], puis remplit le tableau suivant avec les données.

        Args:
            doc: Document Word chargé
            context: Contexte avec les données de tableaux
        """
        logger.info("Remplissage des tableaux dynamiques...")

        # Parcourir les paragraphes pour trouver les marqueurs
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            # Si c'est un marqueur de tableau
            if text.startswith('[TABLE:') and text.endswith(']'):
                table_name = text[7:-1].lower()  # Extraire le nom
                logger.info(f"Marqueur trouvé: {table_name}")

                # Trouver le tableau suivant
                table = self._find_next_table(doc, para_idx)
                if table and table_name in context:
                    data = context[table_name]
                    if isinstance(data, list) and len(data) > 0:
                        self._populate_table(table, data)
                        logger.info(f"✅ Tableau '{table_name}' rempli avec {len(data)} lignes")
                    else:
                        logger.warning(f"⚠️ Données vides pour '{table_name}'")
                else:
                    logger.warning(f"⚠️ Tableau ou données introuvables pour '{table_name}'")

    def _find_next_table(self, doc: Document, para_idx: int):
        """
        Trouve le prochain tableau après un paragraphe donné.

        Args:
            doc: Document Word
            para_idx: Index du paragraphe marqueur

        Returns:
            Le tableau suivant ou None
        """
        # Parcourir les éléments après le paragraphe
        para_element = doc.paragraphs[para_idx]._element
        next_element = para_element.getnext()

        while next_element is not None:
            # Si c'est un tableau
            if next_element.tag.endswith('tbl'):
                # Trouver l'objet Table correspondant
                for table in doc.tables:
                    if table._element == next_element:
                        return table
            next_element = next_element.getnext()

        return None

    def _populate_table(self, table, data: List[Dict[str, Any]]) -> None:
        """
        Remplit un tableau avec des données.

        Les données doivent être dans l'ordre des colonnes du tableau.
        Python 3.7+ garantit l'ordre d'insertion des dicts.

        Args:
            table: Tableau Word
            data: Liste de dictionnaires avec les données (ordre des clés = ordre des colonnes)
        """
        # Chaque dict dans data = une ligne
        for row_data in data:
            # Ajouter une nouvelle ligne
            new_row = table.add_row()

            # Remplir les cellules selon l'ordre des valeurs du dict
            # Note: Python 3.7+ garantit que dict.values() suit l'ordre d'insertion
            for col_idx, value in enumerate(row_data.values()):
                if col_idx < len(new_row.cells):
                    new_row.cells[col_idx].text = str(value)
                else:
                    break

    def _replace_metadata_tags(self, doc: Document, context: Dict[str, Any]) -> None:
        """
        Remplace les balises de métadonnées dans les paragraphes.

        Args:
            doc: Document Word
            context: Contexte avec les métadonnées
        """
        logger.info("Remplacement des balises de métadonnées...")

        replacements = {
            '{{generation_date}}': context.get('generation_date', ''),
            '{{test_start}}': context.get('test_start', ''),
            '{{test_end}}': context.get('test_end', ''),
            '{{turbine_list}}': context.get('turbine_list', ''),
            '{{turbine_count}}': str(context.get('turbine_count', '')),
        }

        # Parcourir tous les paragraphes
        for para in doc.paragraphs:
            for old_tag, new_value in replacements.items():
                if old_tag in para.text:
                    # Remplacer dans chaque run
                    for run in para.runs:
                        if old_tag in run.text:
                            run.text = run.text.replace(old_tag, str(new_value))

        logger.info("✅ Métadonnées remplacées")
